from __future__ import annotations

import argparse
import json
import os
import sys
from copy import deepcopy
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.shared import Inches


DECLARATION_ORDER = [
    ("supplementary_materials", "Supplementary Materials"),
    ("author_contributions", "Author Contributions"),
    ("funding", "Funding"),
    ("institutional_review_board_statement", "Institutional Review Board Statement"),
    ("informed_consent_statement", "Informed Consent Statement"),
    ("data_availability_statement", "Data Availability Statement"),
    ("acknowledgments", "Acknowledgments"),
    ("conflicts_of_interest", "Conflicts of Interest"),
    ("ai_use_disclosure", "AI-Use Disclosure"),
]

NUMERIC_REFERENCE_STYLES = {"vancouver", "ama", "numeric", "numbered", "mdpi"}
REPO_ROOT = Path(__file__).resolve().parents[2]
RULE_INDEX = {
    ("mdpi", "nutrients"): REPO_ROOT / "rules" / "mdpi" / "nutrients.json",
    ("mdpi", "foods"): REPO_ROOT / "rules" / "mdpi" / "foods.json",
    ("mdpi", "journal of clinical medicine"): REPO_ROOT / "rules" / "mdpi" / "jcm.json",
    ("mdpi", "metabolites"): REPO_ROOT / "rules" / "mdpi" / "metabolites.json",
    ("frontiers", "frontiers in nutrition"): REPO_ROOT / "rules" / "frontiers" / "frontiers-in-nutrition.json",
    ("frontiers", "frontiers in pharmacology"): REPO_ROOT / "rules" / "frontiers" / "frontiers-in-pharmacology.json",
    ("plos", "plos one"): REPO_ROOT / "rules" / "plos" / "plos-one.json",
    ("elsevier", "generic elsevier"): REPO_ROOT / "rules" / "elsevier" / "generic-elsevier.json",
}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Generate a Word or LaTeX manuscript skeleton from structured JSON input."
    )
    parser.add_argument("--input", required=True, help="Path to the manuscript JSON input.")
    parser.add_argument(
        "--output-format",
        required=True,
        choices=["word", "latex"],
        help="Output format: word (.docx) or latex (.tex).",
    )
    parser.add_argument("--output", help="Destination file path.")
    parser.add_argument("--style", help="Style profile such as generic or mdpi.")
    parser.add_argument("--journal", help="Journal name override.")
    parser.add_argument("--article-type", help="Article type override.")
    parser.add_argument(
        "--rules",
        help="Optional path to a rule profile JSON file under rules/. If omitted, the script will try to infer one from style and journal.",
    )
    parser.add_argument(
        "--zotero-mode",
        choices=["auto", "required", "disabled"],
        default="auto",
        help="How to handle Zotero bibliography configuration.",
    )
    parser.add_argument(
        "--zotero-bib",
        help="Path to a Zotero-exported .bib file. If omitted, the script checks ZOTERO_BIB_PATH.",
    )
    return parser.parse_args()


def load_payload(path: Path) -> dict[str, Any]:
    with path.open("r", encoding="utf-8") as handle:
        return json.load(handle)


def load_rule_profile(path: Path | None) -> dict[str, Any]:
    if path is None or not path.exists():
        return {}
    with path.open("r", encoding="utf-8") as handle:
        return json.load(handle)


def infer_rule_path(style: str, journal: str) -> Path | None:
    normalized_style = style.strip().lower()
    normalized_journal = journal.strip().lower()
    key = (normalized_style, normalized_journal)
    if key in RULE_INDEX:
        return RULE_INDEX[key]
    slug = normalized_journal.replace("&", "and")
    slug = "".join(ch if ch.isalnum() or ch == " " else " " for ch in slug)
    slug = "-".join(part for part in slug.split() if part)
    style_directory = REPO_ROOT / "rules" / normalized_style
    candidate = style_directory / f"{slug}.json"
    if candidate.exists():
        return candidate
    if style_directory.exists():
        for profile_path in sorted(style_directory.glob("*.json")):
            try:
                profile = load_rule_profile(profile_path)
            except (OSError, json.JSONDecodeError):
                continue
            names = [profile.get("journal"), *(profile.get("aliases") or [])]
            if normalized_journal in {
                str(name).strip().lower() for name in names if name
            }:
                return profile_path
    return None


def validate_generator_article_type(
    rule_profile: dict[str, Any], article_type: str
) -> None:
    policy = rule_profile.get("generator_policy") or {}
    supported = [
        str(item).strip().lower() for item in policy.get("supported_article_types") or []
    ]
    if supported and article_type.strip().lower() not in supported:
        action = policy.get("unsupported_action") or (
            "Use the exact official article-type template and the existing-DOCX "
            "migration workflow instead."
        )
        raise SystemExit(
            f"The structural generator does not support article type '{article_type}' "
            f"for {rule_profile.get('journal') or 'this journal'}. {action}"
        )


def article_sections(style: str, article_type: str, rule_profile: dict[str, Any] | None = None) -> list[dict[str, Any]]:
    normalized = article_type.strip().lower()
    rule_profile = rule_profile or {}
    article_type_sections = rule_profile.get("article_type_sections") or {}
    if normalized in article_type_sections:
        sections = []
        for item in article_type_sections[normalized]:
            if isinstance(item, dict):
                title = str(item.get("title") or "TODO: Section Title")
                suffix = " (optional; retain only when appropriate)." if item.get("optional") else "."
                sections.append(
                    {"title": title, "paragraphs": [f"TODO: Draft the {title.lower()} section{suffix}"]}
                )
            else:
                title = str(item)
                sections.append(
                    {"title": title, "paragraphs": [f"TODO: Draft the {title.lower()} section."]}
                )
        return sections
    if "review" in normalized and "systematic" in normalized:
        titles = [
            "Introduction",
            "Methods",
            "Results",
            "Discussion",
            "Conclusions",
        ]
    elif "review" in normalized:
        titles = [
            "Introduction",
            "Main Sections",
            "Discussion",
            "Conclusions",
        ]
    elif "case report" in normalized:
        titles = [
            "Introduction",
            "Case Presentation",
            "Discussion",
            "Conclusions",
        ]
    else:
        titles = [
            "Introduction",
            "Materials and Methods" if style == "mdpi" else "Methods",
            "Results",
            "Discussion",
            "Conclusions",
        ]
    return [{"title": title, "paragraphs": [f"TODO: Draft the {title.lower()} section."]} for title in titles]


def default_declarations() -> dict[str, str]:
    return {
        "supplementary_materials": "TODO: List supplementary files or confirm not applicable.",
        "author_contributions": "TODO: Describe each author's contribution.",
        "funding": "TODO: State the funding source or write 'This research received no external funding.'",
        "institutional_review_board_statement": "TODO: Provide IRB approval details or confirm not applicable.",
        "informed_consent_statement": "TODO: Provide informed consent details or confirm not applicable.",
        "data_availability_statement": "TODO: Describe data availability or access restrictions.",
        "acknowledgments": "TODO: Add acknowledgments or delete if not applicable.",
        "conflicts_of_interest": "TODO: State conflicts of interest or write 'The authors declare no conflict of interest.'",
        "ai_use_disclosure": "TODO: State whether generative AI or AI-assisted tools require disclosure under the current journal policy.",
    }


def normalize_authors(payload: dict[str, Any]) -> list[dict[str, Any]]:
    authors = payload.get("authors") or []
    if authors:
        return authors
    return [
        {
            "name": "TODO: Author Name",
            "affiliations": [1],
            "email": "TODO: corresponding.author@example.com",
            "corresponding": True,
        }
    ]


def normalize_affiliations(payload: dict[str, Any]) -> list[dict[str, Any]]:
    affiliations = payload.get("affiliations") or []
    if affiliations:
        return affiliations
    return [{"id": 1, "text": "TODO: Author affiliation"}]


def ensure_keyword_placeholders(
    keywords: list[str], keyword_rules: dict[str, Any] | None = None
) -> list[str]:
    keyword_rules = keyword_rules or {}
    items = [str(item) for item in keywords if str(item).strip()]
    target = (
        keyword_rules.get("target_count")
        or keyword_rules.get("min_count")
        or len(items)
        or 3
    )
    while len(items) < target:
        items.append(f"TODO keyword {len(items) + 1}")
    return items


def abstract_labels_for_article_type(
    rule_profile: dict[str, Any], article_type: str
) -> list[str]:
    abstract_rules = rule_profile.get("abstract_rules") or {}
    labels = abstract_rules.get("labels") or []
    if not abstract_rules.get("structured") or not labels:
        return []
    applies_to = [
        str(item).strip().lower() for item in abstract_rules.get("applies_to") or []
    ]
    if applies_to and article_type.strip().lower() not in applies_to:
        return []
    return [str(label) for label in labels]


def structured_abstract_placeholder(
    rule_profile: dict[str, Any], article_type: str
) -> str:
    labels = abstract_labels_for_article_type(rule_profile, article_type)
    if not labels:
        return "TODO: Provide the abstract."
    return " ".join(f"{label}: TODO." for label in labels)


def apply_numbering(sections: list[dict[str, Any]], prefix: str = "") -> list[dict[str, Any]]:
    numbered: list[dict[str, Any]] = []
    for index, section in enumerate(sections, start=1):
        current = f"{prefix}{index}" if not prefix else f"{prefix}.{index}"
        item = deepcopy(section)
        title = str(item.get("title") or "TODO: Section Title")
        if not title.startswith(f"{current}. "):
            item["title"] = f"{current}. {title}"
        if item.get("subsections"):
            item["subsections"] = apply_numbering(item["subsections"], current)
        numbered.append(item)
    return numbered


def normalize_payload(
    payload: dict[str, Any],
    output_format: str,
    default_style: str | None = None,
    default_journal: str | None = None,
    default_article_type: str | None = None,
    rule_profile: dict[str, Any] | None = None,
    style_override: str | None = None,
    journal_override: str | None = None,
    article_type_override: str | None = None,
) -> dict[str, Any]:
    rule_profile = rule_profile or {}
    style = (
        style_override
        or payload.get("style")
        or rule_profile.get("style")
        or default_style
        or "generic"
    ).strip().lower()
    article_type = (
        article_type_override
        or payload.get("article_type")
        or rule_profile.get("default_article_type")
        or default_article_type
        or "Original Research"
    )
    journal = (
        journal_override
        or payload.get("journal")
        or rule_profile.get("journal")
        or default_journal
        or "TODO: Target Journal"
    )
    validate_generator_article_type(rule_profile, str(article_type))
    formatting = rule_profile.get("formatting") or {}
    keyword_rules = rule_profile.get("keyword_rules") or {}
    keywords = ensure_keyword_placeholders(payload.get("keywords") or [], keyword_rules)
    sections = payload.get("sections") or article_sections(style, article_type, rule_profile)
    if formatting.get("section_numbering") and output_format == "word":
        sections = apply_numbering(sections)
    declarations = default_declarations()
    declarations.update(payload.get("declarations") or {})
    normalized = {
        "style": style,
        "output_format": output_format,
        "journal": journal,
        "article_type": article_type,
        "title": payload.get("title") or "TODO: Manuscript Title",
        "running_title": payload.get("running_title") or "",
        "language": payload.get("language") or "English",
        "reference_style": payload.get("reference_style") or ("MDPI" if style == "mdpi" else "Numeric"),
        "citation_mode": rule_profile.get("citation_mode") or payload.get("citation_mode") or "",
        "authors": normalize_authors(payload),
        "affiliations": normalize_affiliations(payload),
        "corresponding_author_note": payload.get("corresponding_author_note") or "",
        "abstract": payload.get("abstract")
        or structured_abstract_placeholder(rule_profile, article_type),
        "keywords": keywords,
        "sections": sections,
        "declarations": declarations,
        "notes": payload.get("notes") or [],
        "rule_profile": rule_profile,
    }
    if rule_profile.get("reference_style") and not payload.get("reference_style"):
        normalized["reference_style"] = rule_profile["reference_style"]
    return normalized


def resolve_output_path(args: argparse.Namespace, payload: dict[str, Any]) -> Path:
    extension = ".docx" if args.output_format == "word" else ".tex"
    if args.output:
        output_path = Path(args.output)
    else:
        stem = payload["journal"].lower().replace(" ", "-").replace("/", "-")
        output_path = Path.cwd() / f"{stem}-manuscript{extension}"
    if output_path.suffix.lower() != extension:
        output_path = output_path.with_suffix(extension)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    return output_path


def resolve_zotero(args: argparse.Namespace, payload: dict[str, Any]) -> dict[str, Any]:
    if args.zotero_mode == "disabled":
        return {
            "enabled": False,
            "path": None,
            "message": "Zotero integration disabled. References section will remain a placeholder.",
        }

    candidate = args.zotero_bib or payload.get("zotero_bib") or os.getenv("ZOTERO_BIB_PATH")
    if candidate:
        bib_path = Path(candidate).expanduser()
        if bib_path.exists():
            return {
                "enabled": True,
                "path": bib_path,
                "message": f"Zotero bibliography source detected: {bib_path}",
            }

    prompt = (
        "Zotero not configured. Provide --zotero-bib <path-to-exported-bib> or set ZOTERO_BIB_PATH. "
        "Leaving the references area as a placeholder."
    )
    if args.zotero_mode == "required":
        raise SystemExit(
            "Zotero not configured. Provide --zotero-bib <path-to-exported-bib> or set ZOTERO_BIB_PATH before running with --zotero-mode required."
        )
    return {"enabled": False, "path": None, "message": prompt}


def get_affiliation_text(affiliations: list[dict[str, Any]], affiliation_id: int) -> str:
    for item in affiliations:
        if item.get("id") == affiliation_id:
            return str(item.get("text") or "")
    return f"TODO: affiliation {affiliation_id}"


def author_line(authors: list[dict[str, Any]]) -> str:
    return ", ".join(str(author.get("name") or "TODO: Author Name") for author in authors)


def corresponding_author_text(authors: list[dict[str, Any]]) -> str:
    corresponding = [author for author in authors if author.get("corresponding")]
    if not corresponding:
        return "TODO: Identify at least one corresponding author and provide a verified email address."
    items = []
    for author in corresponding:
        name = author.get("name") or "TODO: Corresponding Author"
        email = author.get("email") or "TODO: corresponding.author@example.com"
        items.append(f"{name} ({email})")
    return "; ".join(items)


def author_markers(author: dict[str, Any]) -> str:
    markers = [str(item) for item in author.get("affiliations") or []]
    if author.get("equal_contribution"):
        markers.append("†")
    if author.get("corresponding"):
        markers.append("*")
    return ",".join(markers)


def declaration_rows(
    declarations: dict[str, str], rule_profile: dict[str, Any] | None = None
) -> list[tuple[str, str]]:
    rule_profile = rule_profile or {}
    configured = rule_profile.get("declaration_order") or []
    known = {key: title for key, title in DECLARATION_ORDER}
    order: list[tuple[str, str]] = []
    for item in configured:
        if isinstance(item, str) and item in known:
            order.append((item, known[item]))
        elif isinstance(item, dict) and item.get("key"):
            key = str(item["key"])
            order.append((key, str(item.get("label") or known.get(key) or key.replace("_", " ").title())))
    if not order:
        order = DECLARATION_ORDER
    rows = []
    for key, title in order:
        rows.append((title, declarations.get(key) or f"TODO: Complete {title}."))
    return rows


def add_docx_title_page(document: Document, payload: dict[str, Any]) -> None:
    formatting = payload.get("rule_profile", {}).get("formatting") or {}
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(payload["title"])
    run.bold = True
    run.font.size = Pt(formatting.get("title_font_size_pt") or 16)

    if payload["running_title"]:
        running = document.add_paragraph()
        running.alignment = WD_ALIGN_PARAGRAPH.CENTER
        running.add_run(f"Running title: {payload['running_title']}")

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Journal: {payload['journal']}\nArticle type: {payload['article_type']}")

    authors = document.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for index, author in enumerate(payload["authors"]):
        if index:
            authors.add_run(", ")
        authors.add_run(str(author.get("name") or "TODO: Author Name"))
        markers = author_markers(author)
        if markers:
            marker_run = authors.add_run(markers)
            marker_run.font.superscript = True

    for affiliation in payload["affiliations"]:
        para = document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run(f"{affiliation.get('id')}. {affiliation.get('text')}")

    corresponding = document.add_paragraph()
    corresponding.alignment = WD_ALIGN_PARAGRAPH.CENTER
    corresponding.add_run(f"* Correspondence: {corresponding_author_text(payload['authors'])}")

    if any(author.get("equal_contribution") for author in payload["authors"]):
        equal = document.add_paragraph()
        equal.alignment = WD_ALIGN_PARAGRAPH.CENTER
        equal.add_run("† These authors contributed equally to this work.")

    if payload["corresponding_author_note"]:
        document.add_paragraph(payload["corresponding_author_note"])

    if formatting.get("title_page_break", True):
        document.add_page_break()


def add_docx_abstract(document: Document, payload: dict[str, Any]) -> None:
    paragraph = document.add_paragraph()
    text = str(payload["abstract"])
    labels = abstract_labels_for_article_type(
        payload.get("rule_profile", {}), str(payload.get("article_type") or "")
    )
    if not labels:
        paragraph.add_run(text)
        return
    cursor = 0
    for label in labels:
        token = f"{label}:"
        position = text.find(token, cursor)
        if position < 0:
            continue
        if position > cursor:
            paragraph.add_run(text[cursor:position])
        paragraph.add_run(token).bold = True
        cursor = position + len(token)
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def add_docx_sections(document: Document, sections: list[dict[str, Any]], level: int = 1) -> None:
    heading_level = min(level, 9)
    for section in sections:
        document.add_heading(section.get("title") or "TODO: Section Title", level=heading_level)
        for paragraph in section.get("paragraphs") or ["TODO: Add section content."]:
            document.add_paragraph(paragraph)
        if section.get("subsections"):
            add_docx_sections(document, section["subsections"], level + 1)


def generate_docx(payload: dict[str, Any], output_path: Path, zotero: dict[str, Any]) -> None:
    document = Document()
    formatting = payload.get("rule_profile", {}).get("formatting") or {}
    margin_in = formatting.get("margin_in")
    for section in document.sections:
        if formatting.get("page_width_in"):
            section.page_width = Inches(formatting["page_width_in"])
        if formatting.get("page_height_in"):
            section.page_height = Inches(formatting["page_height_in"])
        if margin_in or formatting.get("top_margin_in"):
            section.top_margin = Inches(formatting.get("top_margin_in") or margin_in)
        if margin_in or formatting.get("bottom_margin_in"):
            section.bottom_margin = Inches(formatting.get("bottom_margin_in") or margin_in)
        if margin_in or formatting.get("left_margin_in"):
            section.left_margin = Inches(formatting.get("left_margin_in") or margin_in)
        if margin_in or formatting.get("right_margin_in"):
            section.right_margin = Inches(formatting.get("right_margin_in") or margin_in)
        if formatting.get("continuous_line_numbering"):
            for existing in section._sectPr.findall(qn("w:lnNumType")):
                section._sectPr.remove(existing)
            line_numbers = OxmlElement("w:lnNumType")
            line_numbers.set(qn("w:countBy"), "1")
            line_numbers.set(qn("w:restart"), "continuous")
            line_numbers.set(qn("w:distance"), str(formatting.get("line_number_distance_twips") or 255))
            section._sectPr.insert_element_before(
                line_numbers,
                "w:pgNumType",
                "w:cols",
                "w:formProt",
                "w:vAlign",
                "w:noEndnote",
                "w:titlePg",
                "w:textDirection",
                "w:bidi",
                "w:rtlGutter",
                "w:docGrid",
                "w:printerSettings",
                "w:sectPrChange",
            )
    normal = document.styles["Normal"]
    normal.font.name = formatting.get("font_name") or "Times New Roman"
    normal.font.size = Pt(formatting.get("font_size_pt") or 12)
    if formatting.get("line_spacing"):
        normal.paragraph_format.line_spacing = formatting["line_spacing"]

    add_docx_title_page(document, payload)

    if formatting.get("front_matter_as_heading", True):
        document.add_heading("Abstract", level=1)
    else:
        document.add_paragraph().add_run("Abstract").bold = True
    add_docx_abstract(document, payload)

    keyword_separator = (payload.get("rule_profile", {}).get("keyword_rules") or {}).get("separator") or ", "
    if formatting.get("inline_keywords"):
        keyword_paragraph = document.add_paragraph()
        keyword_paragraph.add_run("Keywords: ").bold = True
        keyword_paragraph.add_run(keyword_separator.join(payload["keywords"]))
    else:
        document.add_heading("Keywords", level=1)
        document.add_paragraph(keyword_separator.join(payload["keywords"]))

    add_docx_sections(document, payload["sections"])

    if formatting.get("declarations_heading", True):
        document.add_heading("Declarations", level=1)
    for title, value in declaration_rows(payload["declarations"], payload.get("rule_profile")):
        para = document.add_paragraph()
        para.add_run(f"{title}: ").bold = True
        para.add_run(value)

    document.add_heading("References", level=1)
    if zotero["enabled"]:
        document.add_paragraph(
            f"{zotero['message']}. Refresh citations and bibliography from Zotero in Word before submission."
        )
    else:
        document.add_paragraph(zotero["message"])
    document.add_paragraph("TODO: Insert or refresh the final bibliography here.")

    document.save(output_path)


def latex_escape(text: str) -> str:
    replacements = {
        "\\": "\\textbackslash{}",
        "&": "\\&",
        "%": "\\%",
        "$": "\\$",
        "#": "\\#",
        "_": "\\_",
        "{": "\\{",
        "}": "\\}",
        "~": "\\textasciitilde{}",
        "^": "\\textasciicircum{}",
    }
    return "".join(replacements.get(character, character) for character in text)


def latex_section_command(level: int) -> str:
    if level <= 1:
        return "section"
    if level == 2:
        return "subsection"
    return "subsubsection"


def build_latex_sections(sections: list[dict[str, Any]], level: int = 1) -> list[str]:
    lines: list[str] = []
    for section in sections:
        command = latex_section_command(level)
        lines.append(f"\\{command}{{{latex_escape(section.get('title') or 'TODO: Section Title')}}}")
        for paragraph in section.get("paragraphs") or ["TODO: Add section content."]:
            lines.append(latex_escape(paragraph))
            lines.append("")
        if section.get("subsections"):
            lines.extend(build_latex_sections(section["subsections"], level + 1))
    return lines


def biblatex_style(reference_style: str, citation_mode: str | None = None) -> str:
    explicit_mode = (citation_mode or "").strip().lower()
    if explicit_mode in {"numeric", "numbered"}:
        return "numeric"
    if explicit_mode in {"author-date", "author date", "authoryear"}:
        return "authoryear"
    normalized = reference_style.strip().lower()
    if any(
        marker in normalized
        for marker in ("non-numeric", "non numeric", "author-date", "author date", "authoryear")
    ):
        return "authoryear"
    tokens = set(normalized.replace("-", " ").split())
    return (
        "numeric"
        if normalized in NUMERIC_REFERENCE_STYLES
        or tokens.intersection({"numeric", "numbered", "vancouver", "ama"})
        else "authoryear"
    )


def latex_author_line(authors: list[dict[str, Any]]) -> str:
    rendered = []
    for author in authors:
        name = latex_escape(str(author.get("name") or "TODO: Author Name"))
        markers = author_markers(author)
        if markers:
            rendered_markers = ",".join(
                "$\\dagger$" if marker == "†" else latex_escape(marker)
                for marker in markers.split(",")
            )
            name += f"\\textsuperscript{{{rendered_markers}}}"
        rendered.append(name)
    return ", ".join(rendered)


def latex_abstract_text(payload: dict[str, Any]) -> str:
    text = str(payload["abstract"])
    labels = abstract_labels_for_article_type(
        payload.get("rule_profile", {}), str(payload.get("article_type") or "")
    )
    if not labels:
        return latex_escape(text)
    parts: list[str] = []
    cursor = 0
    for label in labels:
        token = f"{label}:"
        position = text.find(token, cursor)
        if position < 0:
            continue
        if position > cursor:
            parts.append(latex_escape(text[cursor:position]))
        parts.append(f"\\textbf{{{latex_escape(token)}}}")
        cursor = position + len(token)
    if cursor < len(text):
        parts.append(latex_escape(text[cursor:]))
    return "".join(parts)


def latex_geometry_options(formatting: dict[str, Any]) -> str:
    options = []
    mapping = [
        ("page_width_in", "paperwidth"),
        ("page_height_in", "paperheight"),
        ("top_margin_in", "top"),
        ("bottom_margin_in", "bottom"),
        ("left_margin_in", "left"),
        ("right_margin_in", "right"),
    ]
    for profile_key, latex_key in mapping:
        if formatting.get(profile_key) is not None:
            options.append(f"{latex_key}={formatting[profile_key]}in")
    if not any(formatting.get(key) is not None for key, _ in mapping):
        options.append(f"margin={formatting.get('margin_in') or 1}in")
    return ",".join(options)


def generate_latex(payload: dict[str, Any], output_path: Path, zotero: dict[str, Any]) -> None:
    formatting = payload.get("rule_profile", {}).get("formatting") or {}
    font_size = formatting.get("font_size_pt") or 12
    line_spacing = formatting.get("line_spacing") or 1.5
    lines = [
        f"\\documentclass[{font_size}pt]{{article}}",
        f"\\usepackage[{latex_geometry_options(formatting)}]{{geometry}}",
        "\\usepackage[hidelinks]{hyperref}",
        "\\usepackage{setspace}",
    ]
    if formatting.get("continuous_line_numbering"):
        lines.append("\\usepackage{lineno}")
    if zotero["enabled"]:
        lines.extend(
            [
                f"\\usepackage[backend=biber,style={biblatex_style(payload['reference_style'], payload.get('citation_mode'))}]{{biblatex}}",
                f"\\addbibresource{{{latex_escape(str(zotero['path']))}}}",
            ]
        )
    title_environment = "titlepage" if formatting.get("title_page_break", True) else "center"
    lines.extend(
        [
            "",
            "\\begin{document}",
            f"\\setstretch{{{line_spacing}}}",
            "\\linenumbers" if formatting.get("continuous_line_numbering") else "",
            f"\\begin{{{title_environment}}}",
            "\\centering",
            f"{{\\LARGE \\textbf{{{latex_escape(payload['title'])}}}\\\\[1em]}}",
        ]
    )
    if payload["running_title"]:
        lines.append(f"{{\\large Running title: {latex_escape(payload['running_title'])}\\\\[0.5em]}}")
    lines.append(f"{{\\large Journal: {latex_escape(payload['journal'])}\\\\}}")
    lines.append(f"{{\\large Article type: {latex_escape(payload['article_type'])}\\\\[1em]}}")
    lines.append(f"{{\\large {latex_author_line(payload['authors'])}\\\\[1em]}}")
    for affiliation in payload["affiliations"]:
        lines.append(
            f"{{\\normalsize {affiliation.get('id')}. {latex_escape(str(affiliation.get('text') or ''))}\\\\}}"
        )
    lines.append("")
    lines.append(
        f"{{\\normalsize Corresponding author: {latex_escape(corresponding_author_text(payload['authors']))}\\\\}}"
    )
    if payload["corresponding_author_note"]:
        lines.append(f"{{\\normalsize {latex_escape(payload['corresponding_author_note'])}\\\\}}")
    if any(author.get("equal_contribution") for author in payload["authors"]):
        lines.append("{\\normalsize \\textsuperscript{$\\dagger$} These authors contributed equally to this work.\\\\}")
    if title_environment == "titlepage":
        lines.append("\\vfill")
    lines.extend([f"\\end{{{title_environment}}}", ""])
    lines.extend(["\\begin{abstract}", latex_abstract_text(payload), "\\end{abstract}", ""])
    keyword_separator = (payload.get("rule_profile", {}).get("keyword_rules") or {}).get("separator") or ", "
    lines.append(
        f"\\noindent\\textbf{{Keywords:}} {latex_escape(keyword_separator.join(payload['keywords']))}"
    )
    lines.append("")
    lines.extend(build_latex_sections(payload["sections"]))
    declarations_heading = formatting.get("declarations_heading", True)
    if declarations_heading:
        lines.extend(["\\section*{Declarations}", ""])
    for title, value in declaration_rows(payload["declarations"], payload.get("rule_profile")):
        if declarations_heading:
            lines.append(f"\\subsection*{{{latex_escape(title)}}}")
            lines.append(latex_escape(value))
        else:
            lines.append(
                f"\\noindent\\textbf{{{latex_escape(title)}:}} {latex_escape(value)}"
            )
        lines.append("")
    if zotero["enabled"]:
        lines.append(f"% {latex_escape(zotero['message'])}")
        lines.append("\\printbibliography")
    else:
        lines.extend(
            [
                "\\section*{References}",
                latex_escape(zotero["message"]),
                "",
                "TODO: Insert the final reference list here or configure Zotero before compiling.",
            ]
        )
    lines.extend(["", "\\end{document}", ""])
    output_path.write_text("\n".join(lines), encoding="utf-8")


def main(
    default_style: str | None = None,
    default_journal: str | None = None,
    default_article_type: str | None = None,
) -> int:
    args = parse_args()
    payload = load_payload(Path(args.input))
    style_hint = (args.style or payload.get("style") or default_style or "generic").strip().lower()
    journal_hint = args.journal or payload.get("journal") or default_journal or "TODO: Target Journal"
    if args.rules:
        rule_path = Path(args.rules).expanduser()
        if not rule_path.exists():
            raise SystemExit(f"Rule profile not found: {rule_path}")
    else:
        rule_path = infer_rule_path(style_hint, journal_hint)
        if rule_path is None:
            print(
                f"Warning: no local rule profile matched style '{style_hint}' and journal "
                f"'{journal_hint}'. Generating a generic structural draft.",
                file=sys.stderr,
            )
    rule_profile = load_rule_profile(rule_path)
    normalized = normalize_payload(
        payload,
        output_format=args.output_format,
        default_style=args.style or default_style,
        default_journal=args.journal or default_journal,
        default_article_type=args.article_type or default_article_type,
        rule_profile=rule_profile,
        style_override=args.style,
        journal_override=args.journal,
        article_type_override=args.article_type,
    )
    output_path = resolve_output_path(args, normalized)
    zotero = resolve_zotero(args, payload)
    if args.output_format == "word":
        generate_docx(normalized, output_path, zotero)
    else:
        generate_latex(normalized, output_path, zotero)
    print(f"Generated {args.output_format} manuscript skeleton: {output_path}")
    if rule_path and rule_path.exists():
        print(f"Applied rule profile: {rule_path}")
    print(zotero["message"])
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
