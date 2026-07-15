from __future__ import annotations

import importlib.util
import json
import base64
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
GENERATOR_PATH = ROOT / "formatting-journal" / "scripts" / "generate_manuscript.py"
WRAPPER_PATH = ROOT / "formatting-metabolites" / "scripts" / "generate_manuscript.py"
QA_PATH = ROOT / "formatting-mdpi" / "scripts" / "qa_docx.py"
RULE_PATH = ROOT / "rules" / "mdpi" / "metabolites.json"
FIXTURE_PATH = ROOT / "tests" / "fixtures" / "metabolites-manuscript.json"


def load_module(name: str, path: Path):
    spec = importlib.util.spec_from_file_location(name, path)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"Unable to load module: {path}")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


GENERATOR = load_module("journal_generator_under_test", GENERATOR_PATH)
QA = load_module("mdpi_qa_under_test", QA_PATH)


class MetabolitesGeneratorTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        cls.rules = json.loads(RULE_PATH.read_text(encoding="utf-8"))
        cls.fixture = json.loads(FIXTURE_PATH.read_text(encoding="utf-8"))

    def test_rule_profile_is_inferred(self) -> None:
        inferred = GENERATOR.infer_rule_path("mdpi", "Metabolites")
        self.assertEqual(inferred, RULE_PATH)
        alias = GENERATOR.infer_rule_path("mdpi", "MDPI Metabolites")
        self.assertEqual(alias, RULE_PATH)

    def test_specialized_article_types_are_not_sent_through_generic_imrad(self) -> None:
        with self.assertRaises(SystemExit) as context:
            GENERATOR.normalize_payload(
                {"article_type": "Protocol"},
                output_format="word",
                rule_profile=self.rules,
            )
        self.assertIn("does not support article type 'Protocol'", str(context.exception))
        self.assertIn("official template", str(context.exception))

    def test_cli_overrides_payload_and_profile(self) -> None:
        normalized = GENERATOR.normalize_payload(
            {
                "style": "generic",
                "journal": "Payload Journal",
                "article_type": "Review"
            },
            output_format="word",
            rule_profile=self.rules,
            style_override="mdpi",
            journal_override="Metabolites",
            article_type_override="Article",
        )
        self.assertEqual(normalized["style"], "mdpi")
        self.assertEqual(normalized["journal"], "Metabolites")
        self.assertEqual(normalized["article_type"], "Article")

    def test_structured_abstract_only_applies_to_configured_article_types(self) -> None:
        article = GENERATOR.structured_abstract_placeholder(self.rules, "Article")
        review = GENERATOR.structured_abstract_placeholder(self.rules, "Review")
        self.assertIn("Background/Objectives: TODO.", article)
        self.assertEqual(review, "TODO: Provide the abstract.")
        self.assertEqual(QA.applicable_abstract_labels(self.rules, "Review"), [])

    def test_review_does_not_receive_structured_abstract_word_count_warning(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory) / "review.docx"
            document = Document()
            document.add_paragraph("Abstract")
            document.add_paragraph("A short unstructured review abstract.")
            document.add_paragraph("Keywords: review; metabolism; synthesis")
            document.add_heading("1. Introduction", level=1)
            document.save(output)
            report = QA.audit(output, None, self.rules, "Review")
            issue_codes = {issue["code"] for issue in report["issues"]}
            self.assertNotIn("abstract-word-count", issue_codes)
            self.assertNotIn("abstract-label-missing", issue_codes)

    def test_qa_detects_split_and_misordered_structured_abstract(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory) / "split-abstract.docx"
            document = Document()
            document.add_paragraph("Abstract")
            document.add_paragraph("Background/Objectives: Context. Methods: Method.")
            document.add_paragraph("Conclusions: Conclusion. Results: Result.")
            document.add_paragraph("Keywords: metabolism; biomarkers; cohort")
            document.add_heading("1. Introduction", level=1)
            document.save(output)
            report = QA.audit(output, None, self.rules, "Article")
            issue_codes = {issue["code"] for issue in report["issues"]}
            self.assertIn("abstract-multiple-paragraphs", issue_codes)
            self.assertIn("abstract-label-order", issue_codes)

    def test_source_comparison_detects_package_object_loss(self) -> None:
        one_pixel_png = base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
        )
        with tempfile.TemporaryDirectory() as directory:
            directory_path = Path(directory)
            image_path = directory_path / "pixel.png"
            image_path.write_bytes(one_pixel_png)
            source_path = directory_path / "source.docx"
            output_path = directory_path / "output.docx"

            source = Document()
            source.add_paragraph("Preserved manuscript text.")
            source.sections[0].header.paragraphs[0].text = "Source header"
            source.add_picture(str(image_path))
            source.save(source_path)

            output = Document()
            output.add_paragraph("Preserved manuscript text.")
            output.save(output_path)

            report = QA.audit(output_path, source_path, {}, "Article")
            issue_codes = [issue["code"] for issue in report["issues"]]
            self.assertIn("object-loss", issue_codes)
            self.assertIn("package-part-changed", issue_codes)
            self.assertEqual(report["comparison"]["text_similarity"], 1.0)
            self.assertLess(report["comparison"]["object_deltas"]["media"], 0)
            self.assertLess(report["comparison"]["object_deltas"]["inline_drawings"], 0)

    def test_mdpi_acs_numeric_maps_to_numeric_biblatex(self) -> None:
        self.assertEqual(GENERATOR.biblatex_style("MDPI ACS numeric"), "numeric")
        self.assertEqual(
            GENERATOR.biblatex_style("ambiguous label", "numeric"), "numeric"
        )
        self.assertEqual(GENERATOR.biblatex_style("non-numeric"), "authoryear")

    def test_latex_escape_does_not_reescape_inserted_commands(self) -> None:
        escaped = GENERATOR.latex_escape(r"A_B & C\D {x}")
        self.assertEqual(
            escaped,
            r"A\_B \& C\textbackslash{}D \{x\}",
        )

    def test_missing_corresponding_author_is_not_inferred(self) -> None:
        text = GENERATOR.corresponding_author_text(
            [{"name": "First Author", "email": "first@example.org"}]
        )
        self.assertIn("TODO", text)
        self.assertNotIn("First Author", text)
        self.assertNotIn("first@example.org", text)

    def test_metabolites_wrapper_generates_expected_docx(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory) / "metabolites.docx"
            result = subprocess.run(
                [
                    sys.executable,
                    str(WRAPPER_PATH),
                    "--input",
                    str(FIXTURE_PATH),
                    "--output-format",
                    "word",
                    "--output",
                    str(output),
                    "--zotero-mode",
                    "disabled",
                ],
                cwd=ROOT,
                capture_output=True,
                text=True,
                check=False,
            )
            self.assertEqual(result.returncode, 0, result.stderr)
            self.assertIn(str(RULE_PATH), result.stdout)
            self.assertTrue(output.exists())

            with ZipFile(output) as archive:
                self.assertIsNone(archive.testzip())
                document_xml = archive.read("word/document.xml").decode("utf-8")
                self.assertIn("<w:lnNumType", document_xml)
                self.assertNotIn('w:type="page"', document_xml)
                self.assertLess(
                    document_xml.index("<w:lnNumType"),
                    document_xml.index("<w:cols"),
                )

            document = Document(output)
            section = document.sections[0]
            self.assertAlmostEqual(section.page_width.inches, 8.2677, places=2)
            self.assertAlmostEqual(section.page_height.inches, 11.6929, places=2)
            self.assertAlmostEqual(section.top_margin.inches, 0.9843, places=2)
            self.assertAlmostEqual(section.bottom_margin.inches, 0.6299, places=2)
            self.assertAlmostEqual(section.left_margin.inches, 0.5, places=2)
            self.assertAlmostEqual(section.right_margin.inches, 0.5, places=2)
            self.assertEqual(document.styles["Normal"].font.name, "Palatino Linotype")
            self.assertAlmostEqual(document.styles["Normal"].font.size.pt, 10.0)

            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            for label in ("Background/Objectives:", "Methods:", "Results:", "Conclusions:"):
                self.assertIn(label, text)
            self.assertIn("Keywords: metabolomics; biomarkers; quality assurance", text)
            self.assertIn("1. Introduction", text)
            self.assertIn("5. Conclusions", text)
            self.assertNotIn("\nDeclarations\n", f"\n{text}\n")

            author_paragraph = next(
                paragraph for paragraph in document.paragraphs if "Alex Chen" in paragraph.text
            )
            superscripts = [
                run.text for run in author_paragraph.runs if run.font.superscript
            ]
            self.assertEqual(superscripts, ["1,†,*", "1,2,†,*"])
            self.assertIn("Alex Chen (alex.chen@example.org)", text)
            self.assertIn("Morgan Li (morgan.li@example.org)", text)

            report = QA.audit(output, None, self.rules, "Article")
            self.assertNotEqual(report["status"], "invalid")
            self.assertFalse(
                [issue for issue in report["issues"] if issue["severity"] == "error"]
            )
            issue_codes = {issue["code"] for issue in report["issues"]}
            self.assertIn("mdpi-styles-missing", issue_codes)
            self.assertIn("unresolved-placeholders", issue_codes)
            self.assertIn("abstract-word-count", issue_codes)
            self.assertNotIn("section-missing", issue_codes)

    def test_metabolites_latex_skeleton_preserves_profile_structure(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory) / "metabolites.tex"
            result = subprocess.run(
                [
                    sys.executable,
                    str(WRAPPER_PATH),
                    "--input",
                    str(FIXTURE_PATH),
                    "--output-format",
                    "latex",
                    "--output",
                    str(output),
                    "--zotero-mode",
                    "disabled",
                ],
                cwd=ROOT,
                capture_output=True,
                text=True,
                check=False,
            )
            self.assertEqual(result.returncode, 0, result.stderr)
            latex = output.read_text(encoding="utf-8")
            self.assertIn("paperwidth=8.2677in", latex)
            self.assertIn("top=0.9843in", latex)
            self.assertIn("\\usepackage{lineno}", latex)
            self.assertIn("\\linenumbers", latex)
            self.assertNotIn("\\begin{titlepage}", latex)
            self.assertNotIn("\\section*{Declarations}", latex)
            self.assertIn("Alex Chen\\textsuperscript{1,$\\dagger$,*}", latex)
            self.assertIn("These authors contributed equally", latex)
            self.assertIn("\\textbf{Background/Objectives:}", latex)

    def test_qa_does_not_treat_body_mentions_as_structure(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory) / "false-positive-check.docx"
            document = Document()
            document.styles["Normal"].font.size = Pt(7)
            document.add_paragraph(
                "This sentence mentions Introduction, Materials and Methods, Results, "
                "Discussion, Background/Objectives:, Methods:, Results:, and Conclusions: "
                "without providing those manuscript structures."
            )
            document.add_paragraph("Introduction")
            document.add_paragraph("Keywords: metabolism, biomarkers, cohort")
            table = document.add_table(rows=1, cols=1)
            table.cell(0, 0).text = "Inherited seven-point table text"
            document.save(output)

            report = QA.audit(output, None, self.rules, "Article")
            issue_codes = [issue["code"] for issue in report["issues"]]
            self.assertIn("section-missing", issue_codes)
            self.assertIn("abstract-label-missing", issue_codes)
            self.assertIn("table-font-small", issue_codes)
            self.assertIn("page-layout-mismatch", issue_codes)
            self.assertIn("line-numbering-missing", issue_codes)
            self.assertIn("body-font-mismatch", issue_codes)
            self.assertIn("body-font-size-mismatch", issue_codes)
            self.assertIn("keyword-separator", issue_codes)
            self.assertIn("back-matter-missing", issue_codes)
            self.assertIn("references-missing", issue_codes)

    def test_explicit_missing_rule_profile_fails(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory) / "should-not-exist.docx"
            missing = Path(directory) / "missing-rules.json"
            result = subprocess.run(
                [
                    sys.executable,
                    str(WRAPPER_PATH),
                    "--input",
                    str(FIXTURE_PATH),
                    "--output-format",
                    "word",
                    "--output",
                    str(output),
                    "--rules",
                    str(missing),
                    "--zotero-mode",
                    "disabled",
                ],
                cwd=ROOT,
                capture_output=True,
                text=True,
                check=False,
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("Rule profile not found", result.stderr)
            self.assertFalse(output.exists())

    def test_invalid_docx_fails_qa_cli(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            invalid = Path(directory) / "invalid.docx"
            invalid.write_text("not a zip package", encoding="utf-8")
            result = subprocess.run(
                [
                    sys.executable,
                    str(QA_PATH),
                    "--input",
                    str(invalid),
                    "--rules",
                    str(RULE_PATH),
                    "--fail-on",
                    "error",
                ],
                cwd=ROOT,
                capture_output=True,
                text=True,
                check=False,
            )
            self.assertEqual(result.returncode, 2)
            report = json.loads(result.stdout)
            self.assertEqual(report["status"], "invalid")
            self.assertEqual(report["severity_counts"], {"error": 1})


if __name__ == "__main__":
    unittest.main()
