from __future__ import annotations

import argparse
import hashlib
import io
import json
import posixpath
import re
from collections import Counter
from difflib import SequenceMatcher
from pathlib import Path
from typing import Any
from urllib.parse import unquote
from zipfile import BadZipFile, ZipFile

from docx import Document

try:
    from PIL import Image
except ImportError:  # pragma: no cover - optional detail only
    Image = None


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
EMU_PER_INCH = 914400
PLACEHOLDER_RE = re.compile(
    r"\bTODO\b|please provide|please complete|\[missing\b|\[verify\b|"
    r"\[source needed\]|10\.3390/x{3,}|\bxxx{2,}\b",
    flags=re.IGNORECASE,
)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Audit a DOCX package and optionally compare it with a format-only source file."
    )
    parser.add_argument("--input", required=True, help="DOCX file to audit.")
    parser.add_argument("--source", help="Original DOCX for preservation comparison.")
    parser.add_argument("--rules", help="Journal rule profile JSON.")
    parser.add_argument("--article-type", help="Article type used for rule checks.")
    parser.add_argument("--json-output", help="Optional path for the JSON report.")
    parser.add_argument(
        "--fail-on",
        choices=["error", "warning", "never"],
        default="error",
        help="Exit non-zero for the selected severity threshold.",
    )
    return parser.parse_args()


def load_rules(path: str | None) -> dict[str, Any]:
    if not path:
        return {}
    rule_path = Path(path).expanduser()
    if not rule_path.exists():
        raise SystemExit(f"Rule profile not found: {rule_path}")
    with rule_path.open("r", encoding="utf-8") as handle:
        return json.load(handle)


def add_issue(
    issues: list[dict[str, str]], severity: str, code: str, message: str
) -> None:
    issues.append({"severity": severity, "code": code, "message": message})


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def token_text_similarity(source_text: str, output_text: str) -> float:
    if source_text == output_text:
        return 1.0
    return SequenceMatcher(
        None,
        source_text.split(),
        output_text.split(),
        autojunk=True,
    ).ratio()


def package_text(archive: ZipFile, part: str = "word/document.xml") -> str:
    from xml.etree import ElementTree as ET

    root = ET.fromstring(archive.read(part))
    return normalize_text(" ".join(node.text or "" for node in root.iter(f"{{{W_NS}}}t")))


def source_part_for_relationships(rel_part: str) -> tuple[str | None, str]:
    if rel_part == "_rels/.rels":
        return None, ""
    directory, filename = posixpath.split(rel_part)
    if posixpath.basename(directory) != "_rels" or not filename.endswith(".rels"):
        return None, ""
    source_directory = posixpath.dirname(directory)
    source_name = filename[: -len(".rels")]
    source_part = posixpath.join(source_directory, source_name)
    return source_part, source_directory


def package_relationship_issues(archive: ZipFile) -> list[str]:
    from xml.etree import ElementTree as ET

    names = set(archive.namelist())
    issues: list[str] = []
    relationship_parts = sorted(name for name in names if name.endswith(".rels"))
    if "word/document.xml" in names and "word/_rels/document.xml.rels" not in names:
        issues.append("Missing word/_rels/document.xml.rels")

    for rel_part in relationship_parts:
        source_part, source_directory = source_part_for_relationships(rel_part)
        rel_root = ET.fromstring(archive.read(rel_part))
        relationships: dict[str, tuple[str, bool]] = {}
        for rel in rel_root.findall(f"{{{REL_NS}}}Relationship"):
            rel_id = rel.attrib.get("Id", "")
            target = unquote(rel.attrib.get("Target", "").split("#", 1)[0])
            external = rel.attrib.get("TargetMode") == "External"
            relationships[rel_id] = (target, external)
            if not external and target:
                if target.startswith("/"):
                    resolved = posixpath.normpath(target.lstrip("/"))
                else:
                    resolved = posixpath.normpath(posixpath.join(source_directory, target))
                if resolved not in names:
                    issues.append(f"{rel_part}: {rel_id} -> missing {resolved}")

        if source_part and source_part in names and source_part.endswith(".xml"):
            try:
                source_root = ET.fromstring(archive.read(source_part))
            except ET.ParseError:
                continue
            referenced: set[str] = set()
            for element in source_root.iter():
                for attribute in (
                    f"{{{R_NS}}}id",
                    f"{{{R_NS}}}embed",
                    f"{{{R_NS}}}link",
                ):
                    if element.attrib.get(attribute):
                        referenced.add(element.attrib[attribute])
            for rel_id in sorted(referenced - set(relationships)):
                issues.append(f"{source_part}: unresolved relationship id {rel_id}")
    return issues


def count_xml_features(archive: ZipFile) -> dict[str, int]:
    from xml.etree import ElementTree as ET

    root = ET.fromstring(archive.read("word/document.xml"))
    return {
        "anchored_drawings": len(root.findall(f".//{{{WP_NS}}}anchor")),
        "inline_drawings": len(root.findall(f".//{{{WP_NS}}}inline")),
        "equations": len(root.findall(f".//{{{M_NS}}}oMath")),
        "hyperlinks": len(root.findall(f".//{{{W_NS}}}hyperlink")),
        "simple_fields": len(root.findall(f".//{{{W_NS}}}fldSimple")),
        "field_instructions": len(root.findall(f".//{{{W_NS}}}instrText")),
        "tracked_insertions": len(root.findall(f".//{{{W_NS}}}ins")),
        "tracked_deletions": len(root.findall(f".//{{{W_NS}}}del")),
        "comment_references": len(root.findall(f".//{{{W_NS}}}commentReference")),
        "footnote_references": len(root.findall(f".//{{{W_NS}}}footnoteReference")),
        "endnote_references": len(root.findall(f".//{{{W_NS}}}endnoteReference")),
        "bookmarks": len(root.findall(f".//{{{W_NS}}}bookmarkStart")),
        "content_controls": len(root.findall(f".//{{{W_NS}}}sdt")),
        "line_number_sections": len(root.findall(f".//{{{W_NS}}}lnNumType")),
    }


def image_metrics(document: Document) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    drawing_tags = {f"{{{WP_NS}}}inline", f"{{{WP_NS}}}anchor"}
    drawings = [element for element in document.element.iter() if element.tag in drawing_tags]
    for drawing in drawings:
        blips = drawing.findall(f".//{{{A_NS}}}blip")
        if not blips:
            continue
        index = len(rows) + 1
        extent = drawing.find(f"{{{WP_NS}}}extent")
        width_emu = int(extent.attrib.get("cx", "0")) if extent is not None else 0
        height_emu = int(extent.attrib.get("cy", "0")) if extent is not None else 0
        width_in = float(width_emu) / EMU_PER_INCH
        height_in = float(height_emu) / EMU_PER_INCH
        row: dict[str, Any] = {
            "index": index,
            "layout": "anchored" if drawing.tag == f"{{{WP_NS}}}anchor" else "inline",
            "display_width_in": round(width_in, 3),
            "display_height_in": round(height_in, 3),
        }
        rel_id = None
        rel_id = blips[0].attrib.get(f"{{{R_NS}}}embed") or blips[0].attrib.get(
            f"{{{R_NS}}}link"
        )
        if rel_id:
            rel = document.part.rels.get(rel_id)
            if rel is not None and rel.is_external:
                row["external_target"] = str(rel.target_ref)
            elif rel is not None and hasattr(rel.target_part, "blob"):
                row["part"] = str(rel.target_part.partname)
                if Image is not None:
                    try:
                        with Image.open(io.BytesIO(rel.target_part.blob)) as image:
                            pixels = image.size
                            row["pixels"] = list(pixels)
                            row["format"] = image.format
                            if width_in and height_in:
                                row["effective_dpi"] = round(
                                    min(pixels[0] / width_in, pixels[1] / height_in), 1
                                )
                    except Exception as exc:  # pragma: no cover - uncommon image codec
                        row["image_read_error"] = str(exc)
        rows.append(row)
    return rows


def xml_node_count(archive: ZipFile, part: str, namespace: str, local_name: str) -> int:
    from xml.etree import ElementTree as ET

    if part not in archive.namelist():
        return 0
    root = ET.fromstring(archive.read(part))
    return len(root.findall(f".//{{{namespace}}}{local_name}")) + int(
        root.tag == f"{{{namespace}}}{local_name}"
    )


def grouped_part_fingerprint(archive: ZipFile, names: list[str]) -> str | None:
    selected = sorted(name for name in names if name in archive.namelist())
    if not selected:
        return None
    digest = hashlib.sha256()
    for name in selected:
        digest.update(name.encode("utf-8"))
        digest.update(b"\0")
        digest.update(archive.read(name))
        digest.update(b"\0")
    return digest.hexdigest()


def style_font_size_pt(style: Any) -> float | None:
    seen: set[str] = set()
    current = style
    while current is not None:
        style_id = str(getattr(current, "style_id", ""))
        if style_id in seen:
            break
        seen.add(style_id)
        font = getattr(current, "font", None)
        size = getattr(font, "size", None)
        if size is not None:
            return float(size.pt)
        current = getattr(current, "base_style", None)
    return None


def style_font_name(style: Any) -> str | None:
    seen: set[str] = set()
    current = style
    while current is not None:
        style_id = str(getattr(current, "style_id", ""))
        if style_id in seen:
            break
        seen.add(style_id)
        font = getattr(current, "font", None)
        name = getattr(font, "name", None)
        if name:
            return str(name)
        current = getattr(current, "base_style", None)
    return None


def effective_run_font_size_pt(run: Any, paragraph: Any, document: Document) -> float | None:
    if run.font.size is not None:
        return float(run.font.size.pt)
    character_style = getattr(run, "style", None)
    inherited = style_font_size_pt(character_style)
    if inherited is not None:
        return inherited
    inherited = style_font_size_pt(paragraph.style)
    if inherited is not None:
        return inherited
    return style_font_size_pt(document.styles["Normal"])


def table_style_font_sizes_pt(table: Any) -> list[float]:
    sizes: list[float] = []
    seen: set[str] = set()
    style = getattr(table, "style", None)
    while style is not None:
        style_id = str(getattr(style, "style_id", ""))
        if style_id in seen:
            break
        seen.add(style_id)
        element = getattr(style, "_element", None)
        if element is not None:
            for size_element in element.findall(f".//{{{W_NS}}}rPr/{{{W_NS}}}sz"):
                value = size_element.attrib.get(f"{{{W_NS}}}val")
                if value and value.isdigit():
                    sizes.append(int(value) / 2)
        style = getattr(style, "base_style", None)
    return sizes


def table_font_issues(
    document: Document, minimum_pt: float
) -> list[tuple[str, str]]:
    findings: list[tuple[str, str]] = []
    for table_index, table in enumerate(document.tables, start=1):
        for size_pt in table_style_font_sizes_pt(table):
            if size_pt < minimum_pt:
                findings.append(
                    (
                        "table-font-small",
                        f"Table {table_index} style contains {size_pt:g} pt text formatting.",
                    )
                )
        for row_index, row in enumerate(table.rows, start=1):
            for cell_index, cell in enumerate(row.cells, start=1):
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        size_pt = effective_run_font_size_pt(run, paragraph, document)
                        if not run.text.strip():
                            continue
                        if size_pt is None:
                            findings.append(
                                (
                                    "table-font-unverified",
                                    f"Table {table_index}, row {row_index}, cell {cell_index}: effective font size could not be resolved.",
                                )
                            )
                        elif size_pt < minimum_pt:
                            findings.append(
                                (
                                    "table-font-small",
                                    f"Table {table_index}, row {row_index}, cell {cell_index}: {size_pt:g} pt effective size",
                                )
                            )
    return findings


def section_metrics(document: Document) -> list[dict[str, Any]]:
    rows = []
    for index, section in enumerate(document.sections, start=1):
        line_numbers = section._sectPr.find(f"{{{W_NS}}}lnNumType")
        rows.append(
            {
                "index": index,
                "page_width_in": round(section.page_width.inches, 3),
                "page_height_in": round(section.page_height.inches, 3),
                "top_margin_in": round(section.top_margin.inches, 3),
                "bottom_margin_in": round(section.bottom_margin.inches, 3),
                "left_margin_in": round(section.left_margin.inches, 3),
                "right_margin_in": round(section.right_margin.inches, 3),
                "line_numbering": line_numbers is not None,
                "line_restart": (
                    line_numbers.attrib.get(f"{{{W_NS}}}restart")
                    if line_numbers is not None
                    else None
                ),
                "line_count_by": (
                    line_numbers.attrib.get(f"{{{W_NS}}}countBy")
                    if line_numbers is not None
                    else None
                ),
                "line_distance_twips": (
                    line_numbers.attrib.get(f"{{{W_NS}}}distance")
                    if line_numbers is not None
                    else None
                ),
            }
        )
    return rows


def add_formatting_issues(
    issues: list[dict[str, str]],
    document: Document,
    metrics: dict[str, Any],
    rules: dict[str, Any],
) -> None:
    formatting = rules.get("formatting") or {}
    expectations = {
        "page_width_in": formatting.get("page_width_in"),
        "page_height_in": formatting.get("page_height_in"),
        "top_margin_in": formatting.get("top_margin_in") or formatting.get("margin_in"),
        "bottom_margin_in": formatting.get("bottom_margin_in") or formatting.get("margin_in"),
        "left_margin_in": formatting.get("left_margin_in") or formatting.get("margin_in"),
        "right_margin_in": formatting.get("right_margin_in") or formatting.get("margin_in"),
    }
    tolerance = float(formatting.get("layout_tolerance_in") or 0.03)
    for section in metrics.get("sections_detail") or []:
        deviations = []
        for key, expected in expectations.items():
            if expected is None:
                continue
            actual = section.get(key)
            if actual is None or abs(float(actual) - float(expected)) > tolerance:
                deviations.append(f"{key}={actual} (expected {expected})")
        if deviations:
            add_issue(
                issues,
                "warning",
                "page-layout-mismatch",
                f"Section {section['index']} layout differs from the rule profile: "
                + "; ".join(deviations),
            )

        if formatting.get("continuous_line_numbering"):
            if not section.get("line_numbering"):
                add_issue(
                    issues,
                    "warning",
                    "line-numbering-missing",
                    f"Section {section['index']} has no line-numbering configuration.",
                )
            elif section.get("line_restart") != "continuous":
                add_issue(
                    issues,
                    "warning",
                    "line-numbering-restart",
                    f"Section {section['index']} restarts line numbering as {section.get('line_restart')!r}.",
                )
            if section.get("line_count_by") not in {None, "1"}:
                add_issue(
                    issues,
                    "warning",
                    "line-numbering-interval",
                    f"Section {section['index']} numbers every {section.get('line_count_by')} lines; expected every line.",
                )
            expected_distance = formatting.get("line_number_distance_twips")
            if expected_distance is not None and section.get("line_distance_twips") != str(
                expected_distance
            ):
                add_issue(
                    issues,
                    "warning",
                    "line-numbering-distance",
                    f"Section {section['index']} line-number distance is {section.get('line_distance_twips')!r}; expected {expected_distance} twips.",
                )

    expected_font = formatting.get("font_name")
    expected_size = formatting.get("font_size_pt")
    configured_styles = formatting.get("body_style_names") or [
        "MDPI_3.1_text",
        "MDPI_3.2_text_no_indent",
    ]
    style_names = {style.name for style in document.styles if style.name}
    inspected = [name for name in configured_styles if name in style_names]
    if not inspected:
        inspected = ["Normal"]
    for style_name in inspected:
        style = document.styles[style_name]
        actual_font = style_font_name(style)
        actual_size = style_font_size_pt(style)
        if expected_font and (
            not actual_font or actual_font.casefold() != str(expected_font).casefold()
        ):
            add_issue(
                issues,
                "warning",
                "body-font-mismatch",
                f"Style {style_name!r} uses {actual_font or 'an unresolved/theme font'!r}; expected {expected_font!r}.",
            )
        if expected_size:
            if actual_size is None:
                add_issue(
                    issues,
                    "warning",
                    "body-font-size-unverified",
                    f"Style {style_name!r} has no resolvable font size; expected {expected_size:g} pt.",
                )
            elif abs(actual_size - float(expected_size)) > 0.1:
                add_issue(
                    issues,
                    "warning",
                    "body-font-size-mismatch",
                    f"Style {style_name!r} uses {actual_size:g} pt; expected {expected_size:g} pt.",
                )


def base_metrics(path: Path) -> tuple[dict[str, Any], str, list[str]]:
    document = Document(path)
    with ZipFile(path) as archive:
        bad_member = archive.testzip()
        if bad_member:
            raise BadZipFile(f"Corrupt ZIP member: {bad_member}")
        names = set(archive.namelist())
        visible_text = package_text(archive)
        xml_features = count_xml_features(archive)
        relationship_issues = package_relationship_issues(archive)
        header_parts = sorted(
            name for name in names if re.fullmatch(r"word/header\d+\.xml", name)
        )
        footer_parts = sorted(
            name for name in names if re.fullmatch(r"word/footer\d+\.xml", name)
        )
        custom_xml_parts = sorted(
            name for name in names if name.startswith("customXml/")
        )
        media_parts = sorted(name for name in names if name.startswith("word/media/"))
        chart_parts = sorted(name for name in names if name.startswith("word/charts/"))
        embedding_parts = sorted(
            name for name in names if name.startswith("word/embeddings/")
        )
        package_parts = {
            "media": len(media_parts),
            "charts": len(chart_parts),
            "embeddings": len(embedding_parts),
            "headers": len(header_parts),
            "header_text_nodes": sum(
                xml_node_count(archive, name, W_NS, "t") for name in header_parts
            ),
            "footers": len(footer_parts),
            "footer_text_nodes": sum(
                xml_node_count(archive, name, W_NS, "t") for name in footer_parts
            ),
            "comments": xml_node_count(archive, "word/comments.xml", W_NS, "comment"),
            "footnotes": xml_node_count(archive, "word/footnotes.xml", W_NS, "footnote"),
            "endnotes": xml_node_count(archive, "word/endnotes.xml", W_NS, "endnote"),
            "custom_xml_parts": len(custom_xml_parts),
            "part_group_fingerprints": {
                "media": grouped_part_fingerprint(archive, media_parts),
                "charts": grouped_part_fingerprint(archive, chart_parts),
                "embeddings": grouped_part_fingerprint(archive, embedding_parts),
                "headers": grouped_part_fingerprint(archive, header_parts),
                "footers": grouped_part_fingerprint(archive, footer_parts),
                "comments": grouped_part_fingerprint(archive, ["word/comments.xml"]),
                "footnotes": grouped_part_fingerprint(archive, ["word/footnotes.xml"]),
                "endnotes": grouped_part_fingerprint(archive, ["word/endnotes.xml"]),
                "custom_xml": grouped_part_fingerprint(archive, custom_xml_parts),
            },
        }

    style_histogram = Counter()
    for paragraph in document.paragraphs:
        style_histogram[paragraph.style.name or paragraph.style.style_id] += 1
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    style_histogram[paragraph.style.name or paragraph.style.style_id] += 1

    placeholder_matches = sorted(set(match.group(0) for match in PLACEHOLDER_RE.finditer(visible_text)))
    metrics = {
        "paragraphs": len(document.paragraphs),
        "nonempty_paragraphs": sum(bool(paragraph.text.strip()) for paragraph in document.paragraphs),
        "tables": len(document.tables),
        "sections": len(document.sections),
        "styles": sorted(style.name or style.style_id for style in document.styles),
        "style_histogram": dict(style_histogram.most_common()),
        "sections_detail": section_metrics(document),
        "images": image_metrics(document),
        "placeholders": placeholder_matches,
        **xml_features,
        **package_parts,
    }
    return metrics, visible_text, relationship_issues


def extract_keywords(document: Document) -> list[str]:
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
    for index, text in enumerate(paragraphs):
        if text.lower().startswith("keywords:"):
            value = text.split(":", 1)[1]
            return [item.strip() for item in re.split(r"[;,]", value) if item.strip()]
        if text.lower() == "keywords" and index + 1 < len(paragraphs):
            return [
                item.strip()
                for item in re.split(r"[;,]", paragraphs[index + 1])
                if item.strip()
            ]
    return []


def keyword_text(document: Document) -> str:
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
    for index, text in enumerate(paragraphs):
        if text.casefold().startswith("keywords:"):
            return text.split(":", 1)[1].strip()
        if text.casefold() == "keywords" and index + 1 < len(paragraphs):
            return paragraphs[index + 1]
    return ""


def strip_section_number(text: str) -> str:
    return re.sub(r"^\s*\d+(?:\.\d+)*\.?\s+", "", text).strip()


def is_section_heading(paragraph: Any) -> bool:
    text = paragraph.text.strip()
    style_name = str(getattr(paragraph.style, "name", "")).casefold()
    return (
        style_name.startswith("heading")
        or style_name.startswith("mdpi_2.")
        or bool(re.match(r"^\s*\d+(?:\.\d+)*\.?\s+", text))
    )


def section_title_occurrence(document: Document, title: str) -> tuple[int, str] | None:
    target = title.strip().casefold()
    combined = {"results", "discussion"}
    for index, paragraph in enumerate(document.paragraphs):
        if not is_section_heading(paragraph):
            continue
        candidate = strip_section_number(paragraph.text).casefold()
        if candidate == target:
            return index, paragraph.text.strip()
        if target in combined and candidate in {"results and discussion", "results & discussion"}:
            return index, paragraph.text.strip()
    return None


def section_title_present(document: Document, title: str) -> bool:
    return section_title_occurrence(document, title) is not None


def extract_abstract_paragraphs(
    document: Document, rules: dict[str, Any], article_type: str
) -> list[str]:
    paragraphs = list(document.paragraphs)
    texts = [paragraph.text.strip() for paragraph in paragraphs]
    stop = len(texts)
    for index, text in enumerate(texts):
        if strip_section_number(text).casefold() == "introduction":
            stop = index
            break

    preamble = texts[:stop]
    for index, text in enumerate(preamble):
        lowered = text.casefold()
        inline_abstract = ""
        if lowered == "abstract":
            start = index + 1
        elif lowered.startswith("abstract:"):
            start = index + 1
            inline_abstract = text.split(":", 1)[1].strip()
        else:
            continue
        block = [inline_abstract] if inline_abstract else []
        for candidate in preamble[start:]:
            if candidate.casefold().startswith("keywords:") or candidate.casefold() == "keywords":
                break
            if candidate:
                block.append(candidate)
        return block

    labels = applicable_abstract_labels(rules, article_type)
    if labels:
        first_token = f"{labels[0]}:"
        for paragraph in paragraphs[:stop]:
            style_name = str(getattr(paragraph.style, "name", "")).casefold()
            if "abstract" in style_name and first_token.casefold() in paragraph.text.casefold():
                return [paragraph.text.strip()]
    return []


def extract_abstract_text(
    document: Document, rules: dict[str, Any], article_type: str
) -> str:
    return " ".join(extract_abstract_paragraphs(document, rules, article_type))


def abstract_word_count(text: str, labels: list[str]) -> int:
    cleaned = text
    for label in labels:
        cleaned = re.sub(rf"\b{re.escape(label)}\s*:", " ", cleaned, flags=re.I)
    return len(re.findall(r"\b[A-Za-z0-9]+(?:[’'-][A-Za-z0-9]+)*\b", cleaned))


def required_section_titles(rules: dict[str, Any], article_type: str) -> list[str]:
    sections = (rules.get("article_type_sections") or {}).get(article_type.lower()) or []
    titles = []
    for item in sections:
        if isinstance(item, dict):
            if item.get("optional"):
                continue
            titles.append(str(item.get("title") or ""))
        else:
            titles.append(str(item))
    return [title for title in titles if title]


BACK_MATTER_LABELS = {
    "supplementary_materials": "Supplementary Materials",
    "author_contributions": "Author Contributions",
    "funding": "Funding",
    "institutional_review_board_statement": "Institutional Review Board Statement",
    "informed_consent_statement": "Informed Consent Statement",
    "data_availability_statement": "Data Availability Statement",
    "acknowledgments": "Acknowledgments",
    "conflicts_of_interest": "Conflicts of Interest",
    "ai_use_disclosure": "AI-Use Disclosure",
}


def paragraph_label_position(document: Document, label: str) -> int | None:
    target = label.casefold()
    for index, paragraph in enumerate(document.paragraphs):
        text = strip_section_number(paragraph.text).strip().casefold()
        if text == target or text.startswith(f"{target}:"):
            return index
    return None


def add_back_matter_issues(
    issues: list[dict[str, str]], document: Document, rules: dict[str, Any]
) -> None:
    back_matter_rules = rules.get("back_matter_rules") or []
    for item in back_matter_rules:
        if not isinstance(item, dict) or not item.get("key"):
            continue
        key = str(item["key"])
        label = BACK_MATTER_LABELS.get(key, key.replace("_", " ").title())
        present = paragraph_label_position(document, label) is not None
        if item.get("required") and not present:
            add_issue(
                issues,
                "warning",
                "back-matter-missing",
                f"Required back-matter item not found: {label}",
            )
        elif item.get("required_when") and not present:
            add_issue(
                issues,
                "warning",
                "back-matter-condition-review",
                f"Confirm whether {label} is required: {item['required_when']}.",
            )

    if rules.get("reference_rules") and paragraph_label_position(document, "References") is None:
        add_issue(
            issues,
            "warning",
            "references-missing",
            "References section was not found.",
        )

    configured = rules.get("declaration_order") or []
    ordered_labels = []
    for item in configured:
        if isinstance(item, str):
            ordered_labels.append(BACK_MATTER_LABELS.get(item, item.replace("_", " ").title()))
        elif isinstance(item, dict) and item.get("key"):
            ordered_labels.append(
                str(item.get("label") or BACK_MATTER_LABELS.get(str(item["key"])) or item["key"])
            )
    ordered_labels.append("References")
    present_positions = [
        (label, position)
        for label in ordered_labels
        if (position := paragraph_label_position(document, label)) is not None
    ]
    positions = [position for _, position in present_positions]
    if positions != sorted(positions):
        add_issue(
            issues,
            "warning",
            "back-matter-order",
            "Back-matter items are not in the configured order: "
            + ", ".join(label for label, _ in present_positions),
        )


def applicable_abstract_labels(rules: dict[str, Any], article_type: str) -> list[str]:
    abstract_rules = rules.get("abstract_rules") or {}
    if not abstract_rules.get("structured"):
        return []
    applies_to = [
        str(item).strip().lower() for item in abstract_rules.get("applies_to") or []
    ]
    if applies_to and article_type.strip().lower() not in applies_to:
        return []
    return [str(label) for label in abstract_rules.get("labels") or []]


def audit(path: Path, source: Path | None, rules: dict[str, Any], article_type: str) -> dict[str, Any]:
    issues: list[dict[str, str]] = []
    try:
        metrics, visible_text, relationship_issues = base_metrics(path)
    except Exception as exc:
        return {
            "input": str(path),
            "status": "invalid",
            "issues": [{"severity": "error", "code": "package-invalid", "message": str(exc)}],
            "severity_counts": {"error": 1},
        }

    for item in relationship_issues:
        add_issue(issues, "error", "relationship-invalid", item)

    document = Document(path)
    add_formatting_issues(issues, document, metrics, rules)

    if metrics["anchored_drawings"]:
        add_issue(
            issues,
            "warning",
            "anchored-drawings",
            f"Found {metrics['anchored_drawings']} floating drawing(s); verify cross-platform placement.",
        )
    if metrics["placeholders"]:
        add_issue(
            issues,
            "warning",
            "unresolved-placeholders",
            "Unresolved author or template placeholders remain: " + ", ".join(metrics["placeholders"]),
        )

    mdpi_styles = [name for name in metrics["styles"] if name.startswith("MDPI_")]
    mdpi_style_usage = sum(
        count
        for name, count in metrics["style_histogram"].items()
        if name.startswith("MDPI_")
    )
    if rules.get("publisher") == "MDPI" and not mdpi_styles:
        add_issue(
            issues,
            "warning",
            "mdpi-styles-missing",
            "No MDPI_* styles were found; this is not evidence of official-template matching.",
        )
    elif rules.get("publisher") == "MDPI" and not mdpi_style_usage:
        add_issue(
            issues,
            "warning",
            "mdpi-styles-unused",
            "MDPI_* styles are defined but no paragraph or table content uses them.",
        )

    configured_body_styles = (rules.get("formatting") or {}).get("body_style_names") or []
    if configured_body_styles and mdpi_styles and not any(
        metrics["style_histogram"].get(style_name, 0) for style_name in configured_body_styles
    ):
        add_issue(
            issues,
            "warning",
            "mdpi-body-style-unused",
            "No content uses the configured MDPI body styles: "
            + ", ".join(configured_body_styles),
        )

    required_titles = required_section_titles(rules, article_type)
    section_positions: list[tuple[str, int]] = []
    for title in required_titles:
        occurrence = section_title_occurrence(document, title)
        if occurrence is None:
            add_issue(issues, "warning", "section-missing", f"Required section not found: {title}")
            continue
        position, raw_text = occurrence
        section_positions.append((title, position))
        if (rules.get("formatting") or {}).get("section_numbering") and not re.match(
            r"^\s*\d+(?:\.\d+)*\.?\s+", raw_text
        ):
            add_issue(
                issues,
                "warning",
                "section-numbering-missing",
                f"Required section is not visibly numbered: {title}",
            )
    if [position for _, position in section_positions] != sorted(
        position for _, position in section_positions
    ):
        add_issue(
            issues,
            "warning",
            "section-order",
            "Required sections are not in the configured order: "
            + ", ".join(title for title, _ in section_positions),
        )

    abstract_rules = rules.get("abstract_rules") or {}
    abstract_labels = applicable_abstract_labels(rules, article_type)
    abstract_paragraphs = extract_abstract_paragraphs(document, rules, article_type)
    abstract_text = " ".join(abstract_paragraphs)
    for label in abstract_labels:
        if f"{label}:".casefold() not in abstract_text.casefold():
            add_issue(issues, "warning", "abstract-label-missing", f"Abstract label not found: {label}")

    label_positions = [abstract_text.casefold().find(f"{label}:".casefold()) for label in abstract_labels]
    if label_positions and all(position >= 0 for position in label_positions):
        if label_positions != sorted(label_positions):
            add_issue(
                issues,
                "warning",
                "abstract-label-order",
                "Structured abstract labels are not in the configured order.",
            )

    if abstract_labels and abstract_rules.get("single_paragraph") and len(abstract_paragraphs) > 1:
        add_issue(
            issues,
            "warning",
            "abstract-multiple-paragraphs",
            f"Structured abstract spans {len(abstract_paragraphs)} paragraphs; the profile requires one paragraph.",
        )

    target_words = abstract_rules.get("target_words")
    enforcement = str(abstract_rules.get("word_count_enforcement") or "warning").lower()
    if (
        target_words
        and abstract_labels
        and abstract_text
        and enforcement not in {"off", "none", "disabled"}
    ):
        count = abstract_word_count(abstract_text, abstract_labels)
        tolerance = float(abstract_rules.get("word_count_tolerance_fraction") or 0.2)
        lower = round(float(target_words) * (1 - tolerance))
        upper = round(float(target_words) * (1 + tolerance))
        if count < lower or count > upper:
            severity = "error" if enforcement == "error" else "warning"
            add_issue(
                issues,
                severity,
                "abstract-word-count",
                f"Abstract has approximately {count} word(s); target is about {target_words} ({lower}–{upper} with configured tolerance).",
            )

    keywords = extract_keywords(document)
    raw_keywords = keyword_text(document)
    keyword_rules = rules.get("keyword_rules") or {}
    minimum = keyword_rules.get("min_count")
    maximum = keyword_rules.get("max_count")
    if minimum is not None and len(keywords) < minimum:
        add_issue(issues, "warning", "keywords-too-few", f"Found {len(keywords)} keyword(s); minimum is {minimum}.")
    if maximum is not None and len(keywords) > maximum:
        add_issue(issues, "warning", "keywords-too-many", f"Found {len(keywords)} keyword(s); maximum is {maximum}.")
    separator = str(keyword_rules.get("separator") or "")
    if separator.strip().startswith(";") and raw_keywords and ";" not in raw_keywords and "," in raw_keywords:
        add_issue(
            issues,
            "warning",
            "keyword-separator",
            "Keywords use commas; the rule profile requires semicolon separation.",
        )

    add_back_matter_issues(issues, document, rules)

    figure_rules = rules.get("figure_rules") or {}
    preferred_dpi = figure_rules.get("preferred_min_dpi")
    allowed_formats = {
        str(item).upper() for item in figure_rules.get("allowed_raster_formats") or []
    }
    for image in metrics["images"]:
        if image.get("external_target"):
            add_issue(
                issues,
                "warning",
                "external-image",
                f"Image {image['index']} is externally linked; embed and verify it before submission.",
            )
        image_format = str(image.get("format") or "").upper()
        if allowed_formats and image_format and image_format not in allowed_formats:
            add_issue(
                issues,
                "warning",
                "figure-format",
                f"Image {image['index']} format is {image_format}; allowed raster formats are {', '.join(sorted(allowed_formats))}.",
            )
        if preferred_dpi:
            dpi = image.get("effective_dpi")
            if dpi is None:
                reason = image.get("image_read_error") or (
                    "external link" if image.get("external_target") else "missing pixel or display-size data"
                )
                add_issue(
                    issues,
                    "warning",
                    "figure-resolution-unchecked",
                    f"Image {image['index']} resolution could not be checked: {reason}.",
                )
            elif dpi < preferred_dpi:
                add_issue(
                    issues,
                    "warning",
                    "figure-resolution",
                    f"Image {image['index']} effective resolution is {dpi:g} dpi; preferred minimum is {preferred_dpi} dpi.",
                )

    minimum_table_font = (rules.get("table_rules") or {}).get("minimum_font_size_pt")
    if minimum_table_font:
        for code, finding in table_font_issues(document, float(minimum_table_font)):
            add_issue(issues, "warning", code, finding)

    comparison = None
    if source:
        try:
            source_metrics, source_text, source_rel_issues = base_metrics(source)
            ratio = token_text_similarity(source_text, visible_text)
            keys = [
                "tables",
                "media",
                "charts",
                "embeddings",
                "anchored_drawings",
                "inline_drawings",
                "equations",
                "hyperlinks",
                "simple_fields",
                "field_instructions",
                "tracked_insertions",
                "tracked_deletions",
                "comment_references",
                "footnote_references",
                "endnote_references",
                "bookmarks",
                "content_controls",
                "headers",
                "header_text_nodes",
                "footers",
                "footer_text_nodes",
                "comments",
                "footnotes",
                "endnotes",
                "custom_xml_parts",
            ]
            deltas = {key: metrics[key] - source_metrics[key] for key in keys}
            source_fingerprints = source_metrics.get("part_group_fingerprints") or {}
            output_fingerprints = metrics.get("part_group_fingerprints") or {}
            fingerprint_changes = {
                key: source_fingerprints.get(key) != output_fingerprints.get(key)
                for key in sorted(set(source_fingerprints) | set(output_fingerprints))
            }
            comparison = {
                "source": str(source),
                "text_similarity": round(ratio, 6),
                "object_deltas": deltas,
                "part_group_changes": fingerprint_changes,
                "source_relationship_issues": source_rel_issues,
            }
            if ratio < 0.995:
                add_issue(
                    issues,
                    "warning",
                    "text-changed",
                    f"Normalized body-token similarity to source is {ratio:.4%}; explain all content differences.",
                )
            for key, delta in deltas.items():
                if delta < 0:
                    add_issue(
                        issues,
                        "warning",
                        "object-loss",
                        f"Output contains {-delta} fewer {key} object(s) than the source.",
                    )
            for key, changed in fingerprint_changes.items():
                if changed:
                    add_issue(
                        issues,
                        "warning",
                        "package-part-changed",
                        f"The {key} package-part group changed; verify that this was intentional and relationships remain valid.",
                    )
        except Exception as exc:
            add_issue(issues, "error", "source-compare-failed", str(exc))

    severity_counts = Counter(item["severity"] for item in issues)
    status = "pass"
    if severity_counts["error"]:
        status = "error"
    elif severity_counts["warning"]:
        status = "warning"
    return {
        "input": str(path),
        "journal": rules.get("journal"),
        "article_type": article_type,
        "status": status,
        "metrics": metrics,
        "comparison": comparison,
        "issues": issues,
        "severity_counts": dict(severity_counts),
        "note": "This structural audit does not replace full-page DOCX rendering and visual inspection.",
    }


def main() -> int:
    args = parse_args()
    input_path = Path(args.input).expanduser()
    if not input_path.exists():
        raise SystemExit(f"Input DOCX not found: {input_path}")
    source_path = Path(args.source).expanduser() if args.source else None
    if source_path is not None and not source_path.exists():
        raise SystemExit(f"Source DOCX not found: {source_path}")
    rules = load_rules(args.rules)
    article_type = args.article_type or rules.get("default_article_type") or "Article"
    report = audit(input_path, source_path, rules, str(article_type))
    rendered = json.dumps(report, ensure_ascii=False, indent=2)
    print(rendered)
    if args.json_output:
        output_path = Path(args.json_output).expanduser()
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(rendered + "\n", encoding="utf-8")

    counts = report.get("severity_counts") or {}
    if args.fail_on == "warning" and (counts.get("warning") or counts.get("error")):
        return 2
    if args.fail_on == "error" and counts.get("error"):
        return 2
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
